"""
Servicio para generar contratos DOCX completados.

Reemplaza placeholders en documentos DOCX con valores proporcionados.
"""

from docx import Document
from io import BytesIO
import json
import re
import logging

logger = logging.getLogger(__name__)


class ContratoGenerator:
    """Genera contratos DOCX reemplazando placeholders con valores"""
    
    def __init__(self):
        logger.info("ContratoGenerator inicializado")
    
    def generar(self, contenido_docx: bytes, campos_editables: dict, valores: dict) -> bytes:
        """
        Reemplaza placeholders en el DOCX con los valores proporcionados.
        
        Args:
            contenido_docx: Bytes del archivo DOCX original
            campos_editables: Dict con la estructura de campos {campos: [...]}
            valores: Dict con {campo_id: valor_ingresado}
        
        Returns:
            bytes del DOCX generado
        """
        try:
            # Cargar documento
            doc = Document(BytesIO(contenido_docx))
            
            # Crear mapeo placeholder -> valor
            mapeo = {}
            for campo in campos_editables.get('campos', []):
                campo_id = campo.get('id')
                placeholder = campo.get('placeholder_original', '[___]')
                valor = valores.get(campo_id, '')
                if valor:
                    # Limpiar placeholder para búsqueda (escapar caracteres especiales)
                    placeholder_escaped = re.escape(placeholder)
                    mapeo[placeholder_escaped] = valor
            
            logger.info(f"ContratoGenerator: Reemplazando {len(mapeo)} placeholders")
            
            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                texto_original = paragraph.text
                texto_nuevo = texto_original
                
                for placeholder, valor in mapeo.items():
                    # Usar regex para reemplazar todas las ocurrencias
                    texto_nuevo = re.sub(placeholder, valor, texto_nuevo)
                
                if texto_nuevo != texto_original:
                    # Limpiar párrafo y agregar nuevo texto
                    paragraph.clear()
                    paragraph.add_run(texto_nuevo)
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            texto_original = paragraph.text
                            texto_nuevo = texto_original
                            
                            for placeholder, valor in mapeo.items():
                                texto_nuevo = re.sub(placeholder, valor, texto_nuevo)
                            
                            if texto_nuevo != texto_original:
                                paragraph.clear()
                                paragraph.add_run(texto_nuevo)
            
            # Guardar a bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            resultado = output.read()
            
            logger.info(f"ContratoGenerator: Documento generado ({len(resultado)} bytes)")
            return resultado
            
        except Exception as e:
            logger.error(f"ContratoGenerator: Error generando documento: {e}", exc_info=True)
            raise

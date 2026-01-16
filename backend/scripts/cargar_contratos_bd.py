#!/usr/bin/env python3
"""
Script para cargar los 133 contratos DOCX a la base de datos.

Lee los archivos DOCX de ~/modelos_machado_docx_v2/ y los carga en la tabla `contratos`.

Uso:
    cd /home/brunogandolfo/cfo-inteligente/backend
    python scripts/cargar_contratos_bd.py
    
    # Ver qué haría sin ejecutar:
    python scripts/cargar_contratos_bd.py --dry-run
    
    # Sin confirmación:
    python scripts/cargar_contratos_bd.py --sin-confirmacion
"""

import sys
import os
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional

# Agregar el directorio raíz al path para imports
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
except ImportError:
    print("ERROR: Falta python-docx. Instalar con: pip install python-docx")
    sys.exit(1)

try:
    from sqlalchemy.exc import SQLAlchemyError
except ImportError:
    print("ERROR: Falta sqlalchemy. Instalar con: pip install sqlalchemy")
    sys.exit(1)

try:
    from app.core.database import SessionLocal
    from app.core.config import settings
    from app.models.contrato import Contrato
except ImportError as e:
    print(f"ERROR: No se pudo importar módulos de la app: {e}")
    print("Asegúrate de ejecutar desde el directorio backend/")
    sys.exit(1)


# =============================================================================
# CONFIGURACIÓN
# =============================================================================

INPUT_DIR = Path.home() / "modelos_machado_docx_v2"
CARPETAS_IGNORAR = {"_backup", "_prueba_piloto", "__pycache__"}


# =============================================================================
# FUNCIONES DE EXTRACCIÓN
# =============================================================================

def extraer_texto_docx(docx_path: Path) -> str:
    """
    Extrae texto plano de un archivo DOCX.
    
    Args:
        docx_path: Ruta al archivo DOCX
    
    Returns:
        Texto plano extraído del documento
    """
    try:
        doc = Document(str(docx_path))
        textos = []
        
        # Extraer texto de párrafos
        for paragraph in doc.paragraphs:
            texto = paragraph.text.strip()
            if texto:
                textos.append(texto)
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto = cell.text.strip()
                    if texto:
                        textos.append(texto)
        
        return "\n\n".join(textos)
    
    except Exception as e:
        print(f"    ⚠️  Error extrayendo texto de {docx_path.name}: {e}")
        return ""


def leer_docx_binario(docx_path: Path) -> bytes:
    """
    Lee el contenido binario completo de un archivo DOCX.
    
    Args:
        docx_path: Ruta al archivo DOCX
    
    Returns:
        Contenido binario del archivo
    """
    try:
        with open(docx_path, 'rb') as f:
            return f.read()
    except Exception as e:
        raise IOError(f"No se pudo leer {docx_path}: {e}")


# =============================================================================
# FUNCIONES DE BD
# =============================================================================
# SessionLocal ya está importado de app.core.database
# Se usa directamente como: session = SessionLocal()


def upsert_contrato(
    session,
    titulo: str,
    categoria: str,
    contenido_docx: bytes,
    contenido_texto: str,
    archivo_original: str,
    dry_run: bool = False
) -> tuple[bool, str]:
    """
    Inserta o actualiza un contrato en la BD.
    
    Args:
        session: Sesión de SQLAlchemy
        titulo: Título del contrato
        categoria: Categoría del contrato
        contenido_docx: Contenido binario del DOCX
        contenido_texto: Texto plano extraído
        archivo_original: Nombre del archivo original
        dry_run: Si True, no hace cambios reales
    
    Returns:
        Tuple (exito: bool, accion: str)
        accion puede ser: "insertado", "actualizado", "error"
    """
    try:
        # Buscar contrato existente
        contrato_existente = session.query(Contrato).filter(
            Contrato.titulo == titulo,
            Contrato.categoria == categoria,
            Contrato.deleted_at.is_(None)
        ).first()
        
        if dry_run:
            if contrato_existente:
                return True, "actualizado (dry-run)"
            else:
                return True, "insertado (dry-run)"
        
        if contrato_existente:
            # Actualizar existente
            contrato_existente.contenido_docx = contenido_docx
            contrato_existente.contenido_texto = contenido_texto
            contrato_existente.archivo_original = archivo_original
            contrato_existente.updated_at = datetime.utcnow()
            accion = "actualizado"
        else:
            # Insertar nuevo
            nuevo_contrato = Contrato(
                titulo=titulo,
                categoria=categoria,
                contenido_docx=contenido_docx,
                contenido_texto=contenido_texto,
                archivo_original=archivo_original,
                fuente_original="machado",
                activo=True
            )
            session.add(nuevo_contrato)
            accion = "insertado"
        
        return True, accion
    
    except Exception as e:
        return False, f"error: {str(e)}"


# =============================================================================
# FUNCIÓN PRINCIPAL
# =============================================================================

def procesar_carpeta(carpeta_path: Path, session, dry_run: bool = False) -> dict:
    """
    Procesa todos los DOCX de una carpeta.
    
    Args:
        carpeta_path: Ruta a la carpeta
        session: Sesión de SQLAlchemy
        dry_run: Si True, no hace cambios reales
    
    Returns:
        Dict con estadísticas: {"insertados": 0, "actualizados": 0, "errores": 0}
    """
    categoria = carpeta_path.name
    stats = {"insertados": 0, "actualizados": 0, "errores": 0}
    
    docx_files = sorted(carpeta_path.glob("*.docx"))
    
    if not docx_files:
        return stats
    
    print(f"\n  📁 {categoria} ({len(docx_files)} archivos)")
    
    for docx_path in docx_files:
        # Ignorar archivos temporales
        if docx_path.name.startswith("~$") or "conversion_log" in docx_path.name:
            continue
        
        titulo = docx_path.stem  # Nombre sin extensión
        archivo_original = docx_path.name
        
        try:
            # Leer contenido binario
            contenido_docx = leer_docx_binario(docx_path)
            
            # Extraer texto plano
            contenido_texto = extraer_texto_docx(docx_path)
            
            # Upsert en BD
            exito, accion = upsert_contrato(
                session=session,
                titulo=titulo,
                categoria=categoria,
                contenido_docx=contenido_docx,
                contenido_texto=contenido_texto,
                archivo_original=archivo_original,
                dry_run=dry_run
            )
            
            if exito:
                if "insertado" in accion:
                    stats["insertados"] += 1
                    print(f"    ✅ {titulo[:50]}... → {accion}")
                elif "actualizado" in accion:
                    stats["actualizados"] += 1
                    print(f"    🔄 {titulo[:50]}... → {accion}")
                else:
                    print(f"    ⚠️  {titulo[:50]}... → {accion}")
            else:
                stats["errores"] += 1
                print(f"    ❌ {titulo[:50]}... → {accion}")
        
        except Exception as e:
            stats["errores"] += 1
            print(f"    ❌ {titulo[:50]}... → error: {str(e)}")
    
    # Commit por carpeta (si no es dry-run)
    if not dry_run:
        try:
            session.commit()
        except SQLAlchemyError as e:
            session.rollback()
            print(f"    ⚠️  Error en commit de {categoria}: {e}")
            stats["errores"] += len(docx_files)
    
    return stats


def main():
    """Función principal."""
    parser = argparse.ArgumentParser(
        description="Cargar contratos DOCX a la base de datos",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  python cargar_contratos_bd.py                    # Con confirmación
  python cargar_contratos_bd.py --dry-run           # Ver qué haría
  python cargar_contratos_bd.py --sin-confirmacion  # Sin preguntar
        """
    )
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Solo mostrar qué haría, sin ejecutar cambios'
    )
    parser.add_argument(
        '--sin-confirmacion',
        action='store_true',
        help='Ejecutar sin pedir confirmación'
    )
    
    args = parser.parse_args()
    
    # Banner
    print("=" * 70)
    print("CARGA DE CONTRATOS A BASE DE DATOS")
    print("=" * 70)
    print(f"Input:  {INPUT_DIR}")
    print(f"BD:     {settings.database_url[:50]}...")
    print(f"Fecha:  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()
    
    # Verificar directorio de entrada
    if not INPUT_DIR.exists():
        print(f"ERROR: No existe el directorio {INPUT_DIR}")
        sys.exit(1)
    
    # Buscar carpetas (categorías)
    carpetas = [
        d for d in INPUT_DIR.iterdir()
        if d.is_dir() and d.name not in CARPETAS_IGNORAR
    ]
    carpetas = sorted(carpetas)
    
    if not carpetas:
        print("No se encontraron carpetas para procesar.")
        sys.exit(0)
    
    # Contar archivos totales
    total_archivos = sum(
        len([f for f in carpeta.glob("*.docx") if not f.name.startswith("~$")])
        for carpeta in carpetas
    )
    
    print(f"CARPETAS ENCONTRADAS: {len(carpetas)}")
    for carpeta in carpetas:
        num_archivos = len([f for f in carpeta.glob("*.docx") if not f.name.startswith("~$")])
        print(f"  {carpeta.name}: {num_archivos} archivos")
    print()
    print(f"TOTAL: {total_archivos} archivos DOCX")
    print()
    
    # Dry run
    if args.dry_run:
        print("MODO DRY-RUN: No se realizarán cambios en la BD.")
        print()
        print("Archivos que se procesarían:")
        for carpeta in carpetas:
            docx_files = sorted([f for f in carpeta.glob("*.docx") if not f.name.startswith("~$")])
            for docx in docx_files:
                print(f"  {carpeta.name}/{docx.name}")
        return
    
    # Confirmación
    if not args.sin_confirmacion:
        print("=" * 70)
        print("⚠️  CONFIRMACIÓN REQUERIDA")
        print("=" * 70)
        print(f"Se van a procesar {total_archivos} archivos DOCX.")
        print(f"Los contratos se cargarán/actualizarán en la tabla 'contratos'.")
        print()
        respuesta = input("¿Continuar? [s/N]: ").strip().lower()
        if respuesta not in ['s', 'si', 'sí', 'y', 'yes']:
            print("Operación cancelada.")
            sys.exit(0)
    
    # Crear sesión de BD
    try:
        session = SessionLocal()
    except Exception as e:
        print(f"ERROR: No se pudo conectar a la BD: {e}")
        sys.exit(1)
    
    # Procesar
    print()
    print("PROCESANDO CONTRATOS...")
    print("-" * 70)
    
    stats_globales = {"insertados": 0, "actualizados": 0, "errores": 0}
    
    try:
        for carpeta in carpetas:
            stats_carpeta = procesar_carpeta(carpeta, session, dry_run=False)
            
            # Acumular estadísticas
            for key in stats_globales:
                stats_globales[key] += stats_carpeta[key]
    
    except KeyboardInterrupt:
        print("\n\n⚠️  Interrumpido por el usuario. Haciendo rollback...")
        session.rollback()
        sys.exit(1)
    
    except Exception as e:
        print(f"\n\n❌ Error fatal: {e}")
        session.rollback()
        sys.exit(1)
    
    finally:
        session.close()
    
    # Reporte final
    print()
    print("=" * 70)
    print("REPORTE FINAL")
    print("=" * 70)
    print(f"Total procesados:  {total_archivos}")
    print(f"✅ Insertados:     {stats_globales['insertados']}")
    print(f"🔄 Actualizados:   {stats_globales['actualizados']}")
    print(f"❌ Errores:        {stats_globales['errores']}")
    print()
    print("¡CARGA COMPLETADA!")


if __name__ == "__main__":
    main()

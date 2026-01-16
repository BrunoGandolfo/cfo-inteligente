#!/usr/bin/env python3
"""
CONVERSIÓN LIMPIA DE CONTRATOS PDF A DOCX
=========================================
Convierte los 133 PDFs de modelos de contratos a DOCX limpios usando PyMuPDF.

Características:
- Extrae SOLO texto (sin imágenes de fondo/marcas de agua)
- Elimina todas las referencias a Machado/Giachero
- Agrega placeholders editables
- Aplica formato profesional (Times New Roman 12pt, márgenes 2.5cm)
- Agrega logos de Conexión Consultora según reglas

Uso:
    python convertir_contratos_limpio.py [--dry-run] [--carpeta NOMBRE]
"""

import re
import sys
import argparse
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import List, Optional

# Verificar dependencias
try:
    import fitz  # PyMuPDF
except ImportError:
    print("ERROR: Falta PyMuPDF. Instalar con: pip install PyMuPDF")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: Falta python-docx. Instalar con: pip install python-docx")
    sys.exit(1)


# =============================================================================
# CONFIGURACIÓN
# =============================================================================

# Directorios
INPUT_DIR = Path.home() / "modelos_machado"
OUTPUT_DIR = Path.home() / "modelos_machado_docx_v2"

# Logos
LOGO_CONEXION = Path("/home/brunogandolfo/cfo-inteligente/frontend/public/logo-conexion.png")
LOGO_X = Path("/home/brunogandolfo/cfo-inteligente/backend/app/static/logo_x_real.png")

# Formato
FUENTE_NOMBRE = "Times New Roman"
FUENTE_TAMANO = 12
MARGEN_CM = 2.5
LOGO_CONEXION_ALTURA_CM = 1.5
LOGO_X_ALTURA_CM = 1.0


# =============================================================================
# PATRONES DE LIMPIEZA
# =============================================================================

# Patrones a ELIMINAR del texto
PATRONES_ELIMINAR = [
    # === MACHADO Y VARIANTES ===
    r"Esc\.?\s*Jorge\s+Julio\s+Machado\s+Giachero",
    r"Esc\.?\s*Jorge\s+Machado\s+Giachero",
    r"Esc\.?\s*Jorge\s+Machado",
    r"Jorge\s+Julio\s+Machado\s+Giachero",
    r"Jorge\s+Julio\s+Machado",
    r"Jorge\s+Machado\s+Giachero",
    r"Jorge\s+Machado",
    r"Machado\s+Giachero",
    r"\bMachado\b",
    
    # === GIACHERO SIN MACHADO ===
    r"Escribano\s+Giachero",
    r"Esc\.?\s*Giachero",
    r"\bGiachero\b",
    
    # === ESTUDIO NOTARIAL ===
    r"Estudio\s+Notarial\s+Machado",
    r"www\.estudionotarialmachado\.com[^\s]*",
    r"estudionotarialmachado\.com[^\s]*",
    r"^ESTUDIO\s+NOTARIAL$",  # Líneas que son solo "ESTUDIO NOTARIAL"
    
    # === CONEXIÓN (limpieza de procesamiento anterior) ===
    r"Modelo\s+adaptado\s+por\s+Conexi[oó]n\s+Consultora[^\n]*",
    r"adaptado\s+por\s+Conexi[oó]n[^\n]*",
    r"Conexi[oó]n\s+Consultora",
    r"www\.conexionconsultora\.com\.uy",
    r"\badaptado\b",
    
    # === FUENTE / AUTOR ===
    r"Fuente\s*:[^\n]*",
    r"Autor\s*:[^\n]*",
    r"\bFuente\b(?!\s+de)",  # "Fuente" solo, pero no "Fuente de" (término legal)
]

# Compilar patrones con flags
REGEX_ELIMINAR = [re.compile(p, re.IGNORECASE) for p in PATRONES_ELIMINAR]

# PALABRAS PROHIBIDAS - Si una línea contiene cualquiera de estas, se elimina COMPLETA
PALABRAS_PROHIBIDAS = [
    r'machado',
    r'giachero', 
    r'estudio\s*notarial',
    r'estudionotarialmachado',
    r'www\.estudionotarial',
]
REGEX_LINEA_PROHIBIDA = re.compile('|'.join(PALABRAS_PROHIBIDAS), re.IGNORECASE)


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class ResultadoConversion:
    """Resultado de la conversión de un archivo."""
    archivo: str
    carpeta: str
    exito: bool = False
    paginas_pdf: int = 0
    parrafos_docx: int = 0
    lineas_originales: int = 0
    lineas_eliminadas: int = 0
    lineas_reunidas: int = 0  # Cuántas líneas se unieron en párrafos
    referencias_limpiadas: int = 0
    placeholders_agregados: int = 0
    error: Optional[str] = None
    advertencias: List[str] = field(default_factory=list)


@dataclass
class ReporteGlobal:
    """Reporte global de la conversión."""
    inicio: datetime = field(default_factory=datetime.now)
    fin: Optional[datetime] = None
    total_archivos: int = 0
    exitosos: int = 0
    fallidos: int = 0
    total_paginas: int = 0
    total_referencias_limpiadas: int = 0
    total_lineas_eliminadas: int = 0
    errores: List[str] = field(default_factory=list)
    resultados: List[ResultadoConversion] = field(default_factory=list)


# =============================================================================
# FUNCIONES DE LIMPIEZA
# =============================================================================

def limpiar_texto(texto: str, resultado: ResultadoConversion) -> str:
    """
    Limpia el texto eliminando referencias no deseadas.
    
    Args:
        texto: Texto a limpiar
        resultado: Objeto para registrar estadísticas
    
    Returns:
        Texto limpiado
    """
    if not texto:
        return ""
    
    texto_limpio = texto
    
    # Aplicar patrones de eliminación
    for regex in REGEX_ELIMINAR:
        matches = regex.findall(texto_limpio)
        if matches:
            resultado.referencias_limpiadas += len(matches)
        texto_limpio = regex.sub('', texto_limpio)
    
    return texto_limpio


def es_linea_eliminar(linea: str) -> bool:
    """
    Verifica si una línea debe eliminarse completamente.
    Si contiene Machado, Giachero, Estudio Notarial → eliminar.
    
    Args:
        linea: Línea de texto a verificar
    
    Returns:
        True si la línea debe eliminarse
    """
    return bool(REGEX_LINEA_PROHIBIDA.search(linea))


def convertir_placeholders(texto: str, resultado: ResultadoConversion) -> str:
    """
    Convierte guiones y puntos suspensivos en placeholders editables.
    
    Args:
        texto: Texto a procesar
        resultado: Objeto para registrar estadísticas
    
    Returns:
        Texto con placeholders
    """
    if not texto:
        return ""
    
    texto_nuevo = texto
    
    # Contar antes de reemplazar
    count_guiones = len(re.findall(r'-{3,}', texto_nuevo))
    count_puntos = len(re.findall(r'\.{4,}', texto_nuevo))
    count_ellipsis = len(re.findall(r'…+', texto_nuevo))
    
    # Reemplazar
    texto_nuevo = re.sub(r'-{3,}', '[___]', texto_nuevo)
    texto_nuevo = re.sub(r'\.{4,}', '[...]', texto_nuevo)
    texto_nuevo = re.sub(r'…+', '[...]', texto_nuevo)
    
    resultado.placeholders_agregados += count_guiones + count_puntos + count_ellipsis
    
    return texto_nuevo


def limpiar_espacios(texto: str) -> str:
    """
    Limpia espacios múltiples y líneas vacías excesivas.
    
    Args:
        texto: Texto a limpiar
    
    Returns:
        Texto con espacios normalizados
    """
    if not texto:
        return ""
    
    # Múltiples espacios → uno solo
    texto = re.sub(r'[ \t]+', ' ', texto)
    
    # Múltiples líneas vacías → máximo dos
    texto = re.sub(r'\n\s*\n\s*\n+', '\n\n', texto)
    
    # Espacios al inicio/final de líneas
    lineas = texto.split('\n')
    lineas = [l.strip() for l in lineas]
    
    return '\n'.join(lineas)


# =============================================================================
# REUNIÓN DE PÁRRAFOS
# =============================================================================

def es_titulo(linea: str) -> bool:
    """
    Detecta si una línea es un título (no debe unirse con la siguiente).
    
    Criterios:
    - TODO MAYÚSCULAS y < 80 caracteres
    - Termina en punto o es corta
    - Empieza con número romano (I, II, III, IV, V, etc.)
    """
    linea = linea.strip()
    if not linea:
        return False
    
    # Todo mayúsculas y relativamente corto
    if linea.isupper() and len(linea) < 80:
        return True
    
    # Empieza con número romano seguido de ) o -
    if re.match(r'^[IVXLC]+[\)\-\.\s]', linea):
        return True
    
    # Títulos comunes de contratos
    titulos_comunes = [
        r'^PRIMERO\b', r'^SEGUNDO\b', r'^TERCERO\b', r'^CUARTO\b', r'^QUINTO\b',
        r'^SEXTO\b', r'^SEPTIMO\b', r'^OCTAVO\b', r'^NOVENO\b', r'^DECIMO\b',
        r'^ARTICULO\b', r'^CLAUSULA\b', r'^CAPITULO\b',
    ]
    for patron in titulos_comunes:
        if re.match(patron, linea, re.IGNORECASE):
            return True
    
    return False


def es_item_lista(linea: str) -> bool:
    """
    Detecta si una línea es un ítem de lista (no debe unirse con la anterior).
    
    Criterios:
    - Empieza con número + ) o .
    - Empieza con letra + ) 
    - Empieza con guion o asterisco
    """
    linea = linea.strip()
    if not linea:
        return False
    
    # Número + ) o . (ej: "1)", "1.", "12)")
    if re.match(r'^\d+[\)\.\-]\s*', linea):
        return True
    
    # Letra + ) (ej: "a)", "b)")
    if re.match(r'^[a-zA-Z][\)]\s*', linea):
        return True
    
    # Guion o asterisco al inicio
    if re.match(r'^[\-\*\•]\s+', linea):
        return True
    
    return False


def termina_oracion(linea: str) -> bool:
    """
    Detecta si una línea termina una oración (puntuación final).
    """
    linea = linea.strip()
    if not linea:
        return True  # Línea vacía = fin de párrafo
    
    # Puntuación que indica fin de oración/cláusula
    return linea[-1] in '.;:!?)'


def empieza_continuacion(linea: str) -> bool:
    """
    Detecta si una línea parece ser continuación de la anterior.
    
    Criterios:
    - Empieza con minúscula
    - Empieza con coma o conjunción
    """
    linea = linea.strip()
    if not linea:
        return False
    
    # Empieza con minúscula
    if linea[0].islower():
        return True
    
    # Empieza con coma (continuación de lista)
    if linea[0] == ',':
        return True
    
    return False


def reunir_parrafos(lineas: List[str]) -> List[str]:
    """
    Une líneas fragmentadas en párrafos coherentes.
    
    Los PDFs tienen saltos de línea duros por el ancho de página.
    Esta función detecta cuándo una línea es continuación de la anterior
    y las une en un solo párrafo.
    
    Args:
        lineas: Lista de líneas extraídas del PDF
    
    Returns:
        Lista de párrafos reunidos
    """
    if not lineas:
        return []
    
    parrafos = []
    parrafo_actual = ""
    
    for linea in lineas:
        linea = linea.strip()
        
        # Línea vacía = separador de párrafos
        if not linea:
            if parrafo_actual:
                parrafos.append(parrafo_actual)
                parrafo_actual = ""
            continue
        
        # Si no hay párrafo actual, empezar uno nuevo
        if not parrafo_actual:
            parrafo_actual = linea
            continue
        
        # Decidir si unir o separar
        debe_separar = False
        
        # La línea actual es un título → separar
        if es_titulo(linea):
            debe_separar = True
        
        # La línea actual es un ítem de lista → separar
        elif es_item_lista(linea):
            debe_separar = True
        
        # El párrafo anterior termina en puntuación final → separar
        elif termina_oracion(parrafo_actual):
            debe_separar = True
        
        # Si no hay razón para separar → unir
        if debe_separar:
            # Guardar párrafo anterior y empezar nuevo
            parrafos.append(parrafo_actual)
            parrafo_actual = linea
        else:
            # Unir con espacio
            parrafo_actual = parrafo_actual + " " + linea
    
    # No olvidar el último párrafo
    if parrafo_actual:
        parrafos.append(parrafo_actual)
    
    return parrafos


# =============================================================================
# FUNCIONES DE FORMATO DOCX
# =============================================================================

def configurar_estilos(doc: Document):
    """
    Configura los estilos del documento.
    
    Args:
        doc: Documento Word
    """
    # Estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = FUENTE_NOMBRE
    font.size = Pt(FUENTE_TAMANO)
    
    # Párrafo
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.0


def configurar_margenes(doc: Document):
    """
    Configura los márgenes del documento.
    
    Args:
        doc: Documento Word
    """
    for section in doc.sections:
        section.top_margin = Cm(MARGEN_CM)
        section.bottom_margin = Cm(MARGEN_CM)
        section.left_margin = Cm(MARGEN_CM)
        section.right_margin = Cm(MARGEN_CM)


def configurar_headers(doc: Document, es_documento_una_pagina: bool):
    """
    Configura los headers con logos según las reglas.
    
    Reglas:
    - Documentos de 1 página: Solo logo Conexión centrado
    - Documentos de 2+ páginas: 
      - Página 1: Logo Conexión centrado
      - Páginas 2+: Logo X en esquina superior derecha
    
    Args:
        doc: Documento Word
        es_documento_una_pagina: True si el documento tiene solo 1 página
    """
    for section in doc.sections:
        # Habilitar header diferente para primera página
        section.different_first_page_header_footer = True
        
        # === HEADER PRIMERA PÁGINA: Logo Conexión centrado ===
        first_header = section.first_page_header
        
        # Limpiar contenido existente
        for para in first_header.paragraphs:
            para.clear()
        
        # Agregar logo centrado
        if first_header.paragraphs:
            para = first_header.paragraphs[0]
        else:
            para = first_header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if LOGO_CONEXION.exists():
            try:
                run = para.add_run()
                run.add_picture(str(LOGO_CONEXION), height=Cm(LOGO_CONEXION_ALTURA_CM))
            except Exception as e:
                pass  # Silenciar error de logo
        
        # === HEADER PÁGINAS SIGUIENTES ===
        regular_header = section.header
        
        # Limpiar contenido existente
        for para in regular_header.paragraphs:
            para.clear()
        
        if es_documento_una_pagina:
            # Documento de 1 página: header vacío en páginas siguientes (no hay)
            pass
        else:
            # Documento de 2+ páginas: Logo X a la derecha
            if regular_header.paragraphs:
                para = regular_header.paragraphs[0]
            else:
                para = regular_header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if LOGO_X.exists():
                try:
                    run = para.add_run()
                    run.add_picture(str(LOGO_X), height=Cm(LOGO_X_ALTURA_CM))
                except Exception as e:
                    pass  # Silenciar error de logo
        
        # === LIMPIAR FOOTERS ===
        for para in section.footer.paragraphs:
            para.clear()
        for para in section.first_page_footer.paragraphs:
            para.clear()


# =============================================================================
# FUNCIÓN PRINCIPAL DE CONVERSIÓN
# =============================================================================

def convertir_pdf_a_docx(pdf_path: Path, docx_path: Path) -> ResultadoConversion:
    """
    Convierte un PDF a DOCX limpio.
    
    Args:
        pdf_path: Ruta al archivo PDF
        docx_path: Ruta de salida para el DOCX
    
    Returns:
        ResultadoConversion con estadísticas
    """
    resultado = ResultadoConversion(
        archivo=pdf_path.name,
        carpeta=pdf_path.parent.name
    )
    
    try:
        # Abrir PDF
        pdf = fitz.open(str(pdf_path))
        resultado.paginas_pdf = len(pdf)
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos y márgenes
        configurar_estilos(doc)
        configurar_margenes(doc)
        
        # Recolectar todas las líneas de todas las páginas primero
        todas_las_lineas = []
        
        for num_pagina, page in enumerate(pdf):
            # Extraer SOLO texto (sin imágenes)
            texto_pagina = page.get_text("text")
            
            # Dividir en líneas y procesar
            lineas = texto_pagina.split('\n')
            resultado.lineas_originales += len(lineas)
            
            for linea in lineas:
                # Verificar si es línea a eliminar completamente
                if es_linea_eliminar(linea):
                    resultado.lineas_eliminadas += 1
                    continue
                
                # Limpiar texto
                linea_limpia = limpiar_texto(linea, resultado)
                
                # Convertir placeholders
                linea_limpia = convertir_placeholders(linea_limpia, resultado)
                
                # Solo agregar si tiene contenido
                if linea_limpia.strip():
                    todas_las_lineas.append(linea_limpia.strip())
        
        # REUNIR líneas fragmentadas en párrafos coherentes
        parrafos = reunir_parrafos(todas_las_lineas)
        
        # Calcular cuántas líneas se reunieron
        resultado.lineas_reunidas = len(todas_las_lineas) - len(parrafos)
        
        # Agregar párrafos al documento
        for parrafo in parrafos:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.add_run(parrafo)
            run.font.name = FUENTE_NOMBRE
            run.font.size = Pt(FUENTE_TAMANO)
            
            resultado.parrafos_docx += 1
        
        pdf.close()
        
        # Configurar headers con logos
        es_una_pagina = resultado.paginas_pdf == 1
        configurar_headers(doc, es_una_pagina)
        
        # Crear directorio de salida si no existe
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Guardar documento
        doc.save(str(docx_path))
        
        resultado.exito = True
        
        # Advertencias
        if resultado.parrafos_docx == 0:
            resultado.advertencias.append("Documento vacío después de conversión")
        
        if not LOGO_CONEXION.exists():
            resultado.advertencias.append("Logo Conexión no encontrado")
        
        if not LOGO_X.exists():
            resultado.advertencias.append("Logo X no encontrado")
        
    except Exception as e:
        resultado.error = str(e)
    
    return resultado


# =============================================================================
# FUNCIONES DE REPORTE
# =============================================================================

def imprimir_resultado(resultado: ResultadoConversion, indice: int, total: int):
    """Imprime el resultado de una conversión."""
    status = "✅" if resultado.exito else "❌"
    print(f"  [{indice:3}/{total}] {status} {resultado.carpeta}/{resultado.archivo}")
    
    if resultado.exito:
        detalles = []
        # Mostrar reducción de párrafos (líneas originales → párrafos finales)
        if resultado.lineas_reunidas > 0:
            detalles.append(f"{resultado.lineas_originales}→{resultado.parrafos_docx} párrafos")
        if resultado.lineas_eliminadas > 0:
            detalles.append(f"elim:{resultado.lineas_eliminadas}")
        if resultado.placeholders_agregados > 0:
            detalles.append(f"placeholders:{resultado.placeholders_agregados}")
        
        if detalles:
            print(f"           ({', '.join(detalles)})")
        
        for adv in resultado.advertencias:
            print(f"           ⚠️  {adv}")
    else:
        print(f"           ERROR: {resultado.error}")


def imprimir_reporte_final(reporte: ReporteGlobal):
    """Imprime el reporte final de la conversión."""
    reporte.fin = datetime.now()
    duracion = reporte.fin - reporte.inicio
    
    print()
    print("=" * 70)
    print("REPORTE FINAL DE CONVERSIÓN")
    print("=" * 70)
    print(f"Inicio:          {reporte.inicio.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Fin:             {reporte.fin.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Duración:        {duracion}")
    print()
    print(f"Total archivos:  {reporte.total_archivos}")
    print(f"Exitosos:        {reporte.exitosos} ({100*reporte.exitosos/max(1,reporte.total_archivos):.1f}%)")
    print(f"Fallidos:        {reporte.fallidos}")
    print()
    print(f"Total páginas procesadas:      {reporte.total_paginas}")
    print(f"Referencias limpiadas:         {reporte.total_referencias_limpiadas}")
    print(f"Líneas eliminadas:             {reporte.total_lineas_eliminadas}")
    print()
    print(f"Output:          {OUTPUT_DIR}")
    
    if reporte.errores:
        print()
        print("ERRORES:")
        for error in reporte.errores:
            print(f"  ❌ {error}")
    
    # Resumen por carpeta
    print()
    print("RESUMEN POR CARPETA:")
    print("-" * 50)
    
    carpetas = {}
    for r in reporte.resultados:
        if r.carpeta not in carpetas:
            carpetas[r.carpeta] = {"ok": 0, "error": 0}
        if r.exito:
            carpetas[r.carpeta]["ok"] += 1
        else:
            carpetas[r.carpeta]["error"] += 1
    
    for carpeta, stats in sorted(carpetas.items()):
        status = "✅" if stats["error"] == 0 else "⚠️"
        print(f"  {status} {carpeta}: {stats['ok']} OK, {stats['error']} errores")


def guardar_log(reporte: ReporteGlobal, log_path: Path):
    """Guarda el log de conversión en un archivo."""
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(f"CONVERSIÓN DE CONTRATOS - LOG\n")
        f.write(f"{'=' * 50}\n")
        f.write(f"Fecha: {reporte.inicio.strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Total: {reporte.total_archivos} archivos\n")
        f.write(f"Exitosos: {reporte.exitosos}\n")
        f.write(f"Fallidos: {reporte.fallidos}\n\n")
        
        f.write("DETALLE:\n")
        f.write("-" * 50 + "\n")
        
        for r in reporte.resultados:
            status = "OK" if r.exito else "ERROR"
            f.write(f"{status}: {r.carpeta}/{r.archivo}\n")
            if r.error:
                f.write(f"  Error: {r.error}\n")
            if r.advertencias:
                for adv in r.advertencias:
                    f.write(f"  Advertencia: {adv}\n")


# =============================================================================
# FUNCIÓN MAIN
# =============================================================================

def main():
    """Función principal."""
    parser = argparse.ArgumentParser(
        description="Convierte PDFs de contratos a DOCX limpios",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  python convertir_contratos_limpio.py              # Procesar todos
  python convertir_contratos_limpio.py --dry-run    # Solo mostrar qué se haría
  python convertir_contratos_limpio.py --carpeta poderes  # Solo una carpeta
        """
    )
    parser.add_argument(
        '--dry-run', 
        action='store_true',
        help='Solo mostrar qué archivos se procesarían, sin convertir'
    )
    parser.add_argument(
        '--carpeta',
        type=str,
        help='Procesar solo una carpeta específica'
    )
    parser.add_argument(
        '--sin-confirmacion',
        action='store_true',
        help='Ejecutar sin pedir confirmación'
    )
    
    args = parser.parse_args()
    
    # Banner
    print("=" * 70)
    print("CONVERSIÓN LIMPIA DE CONTRATOS PDF → DOCX")
    print("=" * 70)
    print(f"Input:  {INPUT_DIR}")
    print(f"Output: {OUTPUT_DIR}")
    print(f"Fecha:  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()
    
    # Verificar directorio de entrada
    if not INPUT_DIR.exists():
        print(f"ERROR: No existe el directorio de entrada: {INPUT_DIR}")
        sys.exit(1)
    
    # Verificar logos
    print("VERIFICANDO LOGOS:")
    if LOGO_CONEXION.exists():
        print(f"  ✅ Logo Conexión: {LOGO_CONEXION}")
    else:
        print(f"  ⚠️  Logo Conexión NO encontrado: {LOGO_CONEXION}")
    
    if LOGO_X.exists():
        print(f"  ✅ Logo X: {LOGO_X}")
    else:
        print(f"  ⚠️  Logo X NO encontrado: {LOGO_X}")
    print()
    
    # Buscar PDFs
    if args.carpeta:
        carpeta_path = INPUT_DIR / args.carpeta
        if not carpeta_path.exists():
            print(f"ERROR: No existe la carpeta: {args.carpeta}")
            print(f"Carpetas disponibles: {[d.name for d in INPUT_DIR.iterdir() if d.is_dir()]}")
            sys.exit(1)
        pdfs = sorted(carpeta_path.glob("*.pdf"))
    else:
        pdfs = sorted(INPUT_DIR.rglob("*.pdf"))
    
    if not pdfs:
        print("No se encontraron archivos PDF para procesar.")
        sys.exit(0)
    
    # Mostrar resumen
    print(f"ARCHIVOS A PROCESAR: {len(pdfs)}")
    
    # Agrupar por carpeta
    por_carpeta = {}
    for pdf in pdfs:
        carpeta = pdf.parent.name
        if carpeta not in por_carpeta:
            por_carpeta[carpeta] = []
        por_carpeta[carpeta].append(pdf)
    
    print()
    for carpeta, archivos in sorted(por_carpeta.items()):
        print(f"  {carpeta}: {len(archivos)} archivos")
    print()
    
    # Dry run
    if args.dry_run:
        print("MODO DRY-RUN: No se realizarán cambios.")
        print()
        print("Archivos que se convertirían:")
        for pdf in pdfs:
            docx_name = pdf.stem + ".docx"
            output_path = OUTPUT_DIR / pdf.parent.name / docx_name
            print(f"  {pdf.parent.name}/{pdf.name} → {output_path}")
        return
    
    # Confirmación
    if not args.sin_confirmacion:
        print("=" * 70)
        print("⚠️  CONFIRMACIÓN REQUERIDA")
        print("=" * 70)
        print(f"Se van a convertir {len(pdfs)} archivos PDF a DOCX.")
        print(f"Los archivos se guardarán en: {OUTPUT_DIR}")
        print()
        respuesta = input("¿Continuar? [s/N]: ").strip().lower()
        if respuesta not in ['s', 'si', 'sí', 'y', 'yes']:
            print("Operación cancelada.")
            sys.exit(0)
    
    # Crear directorio de salida
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # Procesar
    print()
    print("PROCESANDO ARCHIVOS...")
    print("-" * 70)
    
    reporte = ReporteGlobal()
    reporte.total_archivos = len(pdfs)
    
    for i, pdf_path in enumerate(pdfs, 1):
        # Calcular ruta de salida
        carpeta_relativa = pdf_path.parent.name
        docx_name = pdf_path.stem + ".docx"
        docx_path = OUTPUT_DIR / carpeta_relativa / docx_name
        
        # Convertir
        resultado = convertir_pdf_a_docx(pdf_path, docx_path)
        reporte.resultados.append(resultado)
        
        # Actualizar estadísticas
        if resultado.exito:
            reporte.exitosos += 1
            reporte.total_paginas += resultado.paginas_pdf
            reporte.total_referencias_limpiadas += resultado.referencias_limpiadas
            reporte.total_lineas_eliminadas += resultado.lineas_eliminadas
        else:
            reporte.fallidos += 1
            reporte.errores.append(f"{resultado.carpeta}/{resultado.archivo}: {resultado.error}")
        
        # Mostrar progreso
        imprimir_resultado(resultado, i, len(pdfs))
    
    # Guardar log
    log_path = OUTPUT_DIR / f"conversion_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    guardar_log(reporte, log_path)
    
    # Reporte final
    imprimir_reporte_final(reporte)
    
    print()
    print(f"Log guardado en: {log_path}")
    print()
    print("¡CONVERSIÓN COMPLETADA!")


if __name__ == "__main__":
    main()
